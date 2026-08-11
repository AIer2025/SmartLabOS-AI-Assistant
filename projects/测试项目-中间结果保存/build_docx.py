# -*- coding: utf-8 -*-
"""
SmartLabOS 售前 WORD 详细设计方案提案生成脚本 —— 项目：测试项目（成都市食品检验研究院 兽残专项）
数据来源：projects/测试项目/{1.html,2.html,微波辅助萃取-模块-URS.html}、references/01-modules 11张确认模块卡 + 02-platforms。
严禁编造，所有技术参数均可溯源至上述文件。
输出：测试项目_SmartLabOS_Presales_提案_20260630175220.docx
"""
import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\TestClaude\SmartLabOS-AI-Assistant\projects\测试项目\测试项目_SmartLabOS_Presales_提案_20260630175220.docx"

BODY_FONT = "宋体"
HEAD_FONT = "黑体"
BRAND = RGBColor(0x0C, 0x5C, 0x8F)
INK = RGBColor(0x0F, 0x1B, 0x2D)
GREY = RGBColor(0x6B, 0x7C, 0x93)
HDR_FILL = "EEF4FA"

doc = Document()

# ---------- 字体与样式 ----------
def set_eastasia(run, font=BODY_FONT):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)
    rfonts.set(qn('w:eastAsia'), font)

normal = doc.styles['Normal']
normal.font.name = BODY_FONT
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(2)

# 标题样式中文字体
for i, sz in [(1, 18), (2, 14), (3, 12), (4, 11)]:
    st = doc.styles['Heading %d' % i]
    st.font.name = HEAD_FONT
    st.element.rPr.rFonts.set(qn('w:eastAsia'), HEAD_FONT)
    st.font.size = Pt(sz)
    st.font.color.rgb = BRAND if i <= 2 else INK
    st.paragraph_format.space_before = Pt(8)
    st.paragraph_format.space_after = Pt(4)

# 页面 A4
sec = doc.sections[0]
sec.page_width = Mm(210); sec.page_height = Mm(297)
sec.top_margin = Mm(22); sec.bottom_margin = Mm(20)
sec.left_margin = Mm(22); sec.right_margin = Mm(22)
CONTENT_W = sec.page_width - sec.left_margin - sec.right_margin  # EMU

# ---------- 基础写作辅助 ----------
def para(text="", size=10.5, bold=False, color=None, align=None, font=BODY_FONT,
         before=0, after=2, italic=False, indent=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if before: p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent is not None:
        p.paragraph_format.first_line_indent = Pt(indent)
    if text != "":
        r = p.add_run(text)
        r.font.size = Pt(size); r.bold = bold; r.italic = italic
        if color is not None: r.font.color.rgb = color
        set_eastasia(r, font)
    return p

def body(text, size=10.5):
    # 正文段落，首行缩进2字
    return para(text, size=size, indent=21, after=4)

def bullet(text, size=10.5, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.4
    if level:
        p.paragraph_format.left_indent = Pt(21 * (level + 1))
    r = p.add_run(text)
    r.font.size = Pt(size); set_eastasia(r)
    return p

def h(level, text):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_eastasia(r, HEAD_FONT)
    return p

def note(label, text, size=9.5):
    """灰底提示块（用单元格模拟）"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.width = CONTENT_W
    set_cell_shading(cell, "F2F6FB")
    set_cell_border(cell, left="0C5C8F", left_sz=18)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if label:
        r = p.add_run(label + "  "); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = BRAND
        set_eastasia(r)
    r2 = p.add_run(text); r2.font.size = Pt(size); set_eastasia(r2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def set_cell_shading(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexfill)
    tcPr.append(shd)

def set_cell_border(cell, **kw):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        col = kw.get(edge);
        if col is None: continue
        e = OxmlElement('w:%s' % edge)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(kw.get(edge + '_sz', 4)))
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), col)
        borders.append(e)
    tcPr.append(borders)

def _set_run(run, size, bold, color):
    run.font.size = Pt(size); run.bold = bold
    if color is not None: run.font.color.rgb = color
    set_eastasia(run)

def table(headers, rows, widths=None, fontsize=9, header_fill=HDR_FILL,
          header_color=INK, caption=None):
    if caption:
        cp = para(caption, size=9.5, bold=True, color=BRAND, before=4, after=2)
    ncol = len(headers)
    t = doc.add_table(rows=1, cols=ncol)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    if widths is None:
        widths = [1] * ncol
    tot = sum(widths)
    col_emu = [int(CONTENT_W * w / tot) for w in widths]
    # header
    hdr = t.rows[0].cells
    for j, htext in enumerate(headers):
        hdr[j].width = col_emu[j]
        set_cell_shading(hdr[j], header_fill)
        hdr[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = hdr[j].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.15
        _set_run(p.add_run(str(htext)), fontsize, True, header_color)
    # body
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].width = col_emu[j]
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            lines = str(val).split("\n")
            for k, line in enumerate(lines):
                p = cells[j].paragraphs[0] if k == 0 else cells[j].add_paragraph()
                p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.15
                _set_run(p.add_run(line), fontsize, False, None)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def add_field(paragraph, instr):
    r = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin'); r._r.append(b)
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr; r._r.append(it)
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate'); r._r.append(s)
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end'); r._r.append(e)

# ====================================================================
# 封面
# ====================================================================
for _ in range(2):
    doc.add_paragraph()
para("百泉聚兴（北京）科技有限公司", size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("SmartLabOS 智慧实验室自动化系统", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
para("成都市食品检验研究院", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
para("兽药残留专项自动化前处理检测系统", size=20, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, font=HEAD_FONT, after=6)
para("详细设计方案", size=26, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, font=HEAD_FONT, after=8)
para("（售前技术提案 · Pre-sales Technical Proposal）", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)

cover_t = doc.add_table(rows=0, cols=2)
cover_t.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_info = [
    ("项目名称", "测试项目 · 兽残专项自动化检测方案"),
    ("客户单位", "成都市食品检验研究院"),
    ("文档类型", "详细设计方案（详细设计 / 技术评审阶段）"),
    ("覆盖标准", "GB 31658.17—2021、GB/T 21316—2007"),
    ("版本号 / 状态", "V1.0 / 售前提案（待客户技术确认）"),
    ("编制单位", "百泉聚兴（北京）科技有限公司"),
    ("编制日期", "2026 年 6 月 30 日"),
    ("联系人", "请联系销售团队（报价及商务事宜）"),
]
for k, v in cover_info:
    cells = cover_t.add_row().cells
    cells[0].width = Emu(int(CONTENT_W*0.30)); cells[1].width = Emu(int(CONTENT_W*0.55))
    p0 = cells[0].paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run(p0.add_run(k), 10.5, True, BRAND)
    p1 = cells[1].paragraphs[0]
    _set_run(p1.add_run(v), 10.5, False, None)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ====================================================================
# 配套内部文档索引 + 目录
# ====================================================================
h(1, "配套内部文档索引")
para("本详细设计方案基于以下已生成的售前技术文档整合编制，全部技术参数可溯源至下列来源文件与 SmartLabOS 知识库。", size=10)
table(
    ["#", "文件 / 来源", "内容", "对应标准 / 说明"],
    [
        ["1", "projects/测试项目/1.html", "方案一 · 自动化前处理解决方案（HTML 提案）", "GB 31658.17—2021（四环素/磺胺/喹诺酮，41 种）"],
        ["2", "projects/测试项目/2.html", "方案二 · 自动化前处理解决方案（HTML 提案）", "GB/T 21316—2007（磺胺类，23 种）"],
        ["3", "projects/测试项目/微波辅助萃取-模块-URS.html", "新建模块需求规格（URS）", "GB/T 21316 §7.1 微波辐照缺口（MOD-WB-001 拟建）"],
        ["4", "projects/测试项目/模块选定-推荐.json", "客户已确认 11 模块选型清单", "模块选定 → 模块确认 结果"],
        ["5", "projects/测试项目/Summary-输出总结.md", "前一阶段输出总结", "节拍 / 选型 / 偏差汇总"],
        ["6", "references/01-modules、02-platforms、04-solutions", "模块卡 / 平台规格 / 解决方案范式知识库", "全部硬件技术参数唯一来源"],
        ["7", "references/_templates/02-…详细设计方案版.md", "本文档内容范围标准（提纲与勾选项）", "编制依据"],
    ],
    widths=[0.5, 3.2, 3.6, 3.7], fontsize=8.5)

h(1, "目录")
para("（在 Microsoft Word 中打开后，按 Ctrl+A 全选再按 F9，或右键『更新域』即可生成带页码的目录。）", size=9, color=GREY)
toc_p = doc.add_paragraph()
add_field(toc_p, r'TOC \o "1-3" \h \z \u')
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ====================================================================
# 1 项目概述
# ====================================================================
h(1, "1  项目概述")

h(2, "1.1  客户与业务背景")
body("成都市食品检验研究院隶属成都市市场监督管理局，为正处级公益二类事业单位，承担食品、食用农产品等的监督检验、风险监测与技术研究，是国家食品复检机构，牵头食用农产品抽检监测，并开展多兽药残留检测技术攻关。")
body("该院开展食品中硝基呋喃、氯霉素、磺胺类、沙星类（喹诺酮类）四大重点兽残检测。四类兽残均属痕量残留、基质干扰极强，前处理是核心控险环节，对回收率、精密度与结果有效性具有决定性影响。本项目以多品类兽残高通量大筛查为基础，深度整合自动化前处理与上机检测全流程，覆盖样品自动均质、萃取、净化、浓缩、定容、进样等关键环节，全程替代人工繁琐操作。")

h(2, "1.2  检测对象与目标物")
table(
    ["检测标准", "检测项目 / 目标物", "目标物数量", "适用基质"],
    [
        ["GB 31658.17—2021", "四环素类（4）+ 磺胺类（21）+ 喹诺酮（沙星）类（16）", "41 种", "牛、羊、猪、鸡的肌肉 / 肝脏 / 肾脏"],
        ["GB/T 21316—2007", "磺胺脒、甲氧苄啶、磺胺嘧啶等磺胺类", "23 种", "肝、肾、肌肉、水产品（鱼、虾）、牛奶等"],
        ["（专项总览）", "硝基呋喃、氯霉素、磺胺类、沙星类四大重点兽残", "—", "各类动物源性食品基质"],
    ],
    widths=[2.2, 4.6, 1.3, 2.9], fontsize=9)
para("目标物理化特性：四类兽残均为痕量残留（检出限低至 µg/kg 级），基质干扰极强（高脂、高蛋白），部分目标物热敏感、易挥发 / 降解，对均质均一性、净化除杂、控温浓缩、定容精度要求高。", size=9.5, color=GREY)

h(2, "1.3  现状痛点深度分析")
para("人工前处理各环节难点逐项剖析（对应客户提出的 6 项挑战）：", size=10)
pain = [
    ("挑战-1 全程手工容错率极低", "依赖人员熟练度，操作步骤繁琐，个体差异大，难以保证批间一致性。"),
    ("挑战-2 均质不均 / 基质分层", "高脂高蛋白基质易分层，痕量兽残取样偏差大，重复性差，直接影响取样代表性。"),
    ("挑战-3 萃取乳化", "萃取极易乳化，分层困难，目标物在乳化层中损失，回收率下降。"),
    ("挑战-4 SPE 操作失控", "固相萃取柱活化、平衡、流速人工难把控，易柱干、堵塞、填料吸附，回收率波动大。"),
    ("挑战-5 氮吹 / 浓缩 / 定容", "控温控速难，热敏感物易分解、目标物易挥发，蒸干终点难判断，定容精度不足。"),
    ("挑战-6 基质杂质与离子抑制", "基质杂质难除引发严重离子抑制，人工操作个体差异大，直接影响检测回收率、精准度与结果有效性。"),
]
for k, v in pain:
    bullet("%s —— %s" % (k, v), size=9.5)
para("上述痛点的共性机理：人工操作的不可控变量（手法、流速、温度、时间）直接转化为回收率波动与基质效应，最终影响痕量定量结果的准确性与有效性。本方案以参数固化、流程闭环、模块化自动执行从根因消除人为差异。", size=9.5, color=GREY)

h(2, "1.4  项目场地概述")
note("说明（信息待补充）",
     "客户需求资料中暂未提供改造前用途、楼层 / 房间、场地长×宽 / 面积 / 层高及隔断打通范围等场地信息。"
     "依据数据可溯源原则，本节不作场地尺寸编造，列出待客户补充清单，并先行给出由设备规格反推的硬性场地前提条件，供现场勘测与隔断设计依据。")
table(
    ["场地要素", "状态 / 设计前提", "依据"],
    [
        ["改造前用途 / 楼层 / 房间", "待客户提供", "现场勘测确认"],
        ["场地长×宽 / 面积 / 层高", "待客户提供；设备高 2055 mm，要求层高 ≥ 3 m", "PLT-800/1200/1400 高 2055 mm（02-platforms）"],
        ["地面承重", "要求 ≥ 500 kg/㎡", "离心模块 MOD-LX-001 自重 201 kg、浓缩 50 kg 等"],
        ["设备占地（净）", "方案一 6×PLT-800 ≈ 4.08 ㎡；含通道按 2.5~3× 预留 ≈ 10–12 ㎡", "1.html §10 公用工程"],
        ["隔断打通 / 改造范围", "待现场勘测后于阶段一细化", "—"],
    ],
    widths=[2.7, 5.0, 2.3], fontsize=9)

# ====================================================================
# 2 自动化需求与标准适配
# ====================================================================
h(1, "2  自动化需求与标准适配")

h(2, "2.1  设计依据国标清单（现行有效）")
table(
    ["标准号", "标准全称", "状态 / 实施"],
    [
        ["GB 31658.17—2021", "《食品安全国家标准 动物性食品中四环素类、磺胺类和喹诺酮类药物残留量的测定 液相色谱-串联质谱法》", "强制性国标，2022-02-01 实施"],
        ["GB/T 21316—2007", "《动物源性食品中磺胺类药物残留量的测定 液相色谱-质谱/质谱法》", "推荐性国标，2008-04-01 实施"],
        ["GB/T 6682—2008", "《分析实验室用水规格和试验方法》", "纯水 / 上水水质依据"],
    ],
    widths=[2.3, 6.7, 2.0], fontsize=9)

h(2, "2.2  检测标准与设备适配表")
table(
    ["检测方法号", "检测项目", "前处理系统（SmartLabOS）", "检测仪器"],
    [
        ["GB 31658.17—2021", "四环素/磺胺/喹诺酮 41 种", "缓冲液提取 + HLB 固相萃取净化 + 真空离心浓缩 + 定容 + 0.22 µm 过滤（方案一，6 工站）", "LC-MS/MS（液相色谱-串联质谱）"],
        ["GB/T 21316—2007", "磺胺类 23 种", "C₁₈ 研磨分散 + 乙腈微波辅助提取 + 正己烷液液分配净化 + 旋蒸/氮吹浓缩 + 定容 + 过滤（方案二，6 工站）", "LC-MS/MS（液相色谱-质谱/质谱）"],
    ],
    widths=[2.0, 2.2, 5.3, 1.9], fontsize=9)
para("注：两标准前处理路线不同（方案一用 SPE 净化；方案二用微波辅助提取 + 液液分配净化），通过模块分时复用与平台可重构实现一套自动化产线对两类检项的适配。", size=9, color=GREY)

h(2, "2.3  自动化需求统计（通量 / 周期目标）")
table(
    ["方案", "对应标准", "单样全流程串行节拍", "产线瓶颈站节拍", "说明"],
    [
        ["方案一", "GB 31658.17—2021", "3037 s ≈ 50.6 min/样", "WS-02 超声·离心 720 s", "5 处理工站均值 589 s，最接近 600 s 基准"],
        ["方案二", "GB/T 21316—2007", "2737 s ≈ 45.6 min/样", "WS-05 超声·过滤 630 s", "含新建微波模块 MOD-WB-001（120 s/循环）"],
    ],
    widths=[1.2, 2.5, 2.6, 2.4, 2.3], fontsize=9)
para("耗时口径：每平台处理耗时 = Σ 该平台所载模块（功能码单循环耗时 cycle_time_sec + 自动上下料 60 s/模块）；解决方案总耗时 = Σ 各平台耗时。稳态流水生产时各工站可重叠，实际通量由瓶颈站决定，远高于串行单样节拍（详见第 6 章）。", size=9, color=GREY)

h(2, "2.4  差异化检项的系统适配与调度")
bullet("同一套离心管自动化岛（PLT 系列平台 + 11 确认模块）通过软件重构工序，分别加载方案一 / 方案二的流程模板，实现两类标准在同一硬件资产上的切换运行。", size=9.5)
bullet("方案一独有模块：MOD-SPE-004（HLB 固相萃取净化）；方案二独有模块：MOD-WB-001（微波辅助提取，URS 拟建）。其余 9 个模块两方案共用。", size=9.5)
bullet("由 SmartLabOS 智慧实验室软件统一调度，按检项分配工站时序与并行资源，兼顾高通量筛查与靶向定量需求。", size=9.5)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ====================================================================
# 3 总体设计
# ====================================================================
h(1, "3  总体设计")

h(2, "3.1  设计依据与原则")
for t in [
    "模块化：以标准模块位（含 MODBUS TCP 统一总线）为单元，按工序自由组合，便于扩展与重构；",
    "智能物流联动：XYZ 机械臂 + 平台托盘缓存实现工站间样品自动流转，全自动上下料；",
    "数字化管理：智慧实验室软件统一调度、全流程数据采集与样品全生命周期追溯；",
    "可扩展：平台预留模块位与托盘位，瓶颈工站可并行扩容或升级 PLT-1200/1400。",
]:
    bullet(t, size=9.5)
para("业务闭环范围：采样 → 制样 → 分样 → 前处理（均质/提取/净化/浓缩/定容/过滤）→ 检测（LC-MS/MS）→ 出具报告。本方案自动化覆盖范围自『均质制样』起至『待测液冷藏 / 过滤进样』止；称样（1 g / 2 g，感量 0.01 g 天平）在客户已确认 11 模块中无称量模块，按行业惯例由人工称样后上样，已诚实标注。", size=9.5)

h(2, "3.2  整体布局设计")
para("平台 / 区域划分（方案一，6 座工作站，按『每平台处理耗时 ≈ 600 s』均衡拆分）：", size=10)
table(
    ["工作站", "区域职能", "搭载模块", "平台", "处理耗时"],
    [
        ["WS-01 制样提取", "均质 + 加缓冲液 + 涡旋", "MOD-JZ-002 / JY-010 / WX-004", "PLT-800", "566 s"],
        ["WS-02 超声·离心", "超声辅助提取 + 冷冻离心（瓶颈）", "MOD-CS-004 / LX-001", "PLT-800", "720 s"],
        ["WS-03 转移·SPE 净化", "提取液合并 + HLB 固相萃取", "MOD-YY-004 / SPE-004", "PLT-800", "520 s"],
        ["WS-04 浓缩", "45 ℃ 真空离心浓缩", "MOD-NS-001", "PLT-800", "540 s"],
        ["WS-05 定容·过滤", "精密定容 + 0.22 µm 过滤", "MOD-DR-003 / GL-001", "PLT-800", "600 s"],
        ["WS-06 冷链存储缓存", "0–10 ℃ 待测液冷藏 / 进出料缓冲（并行）", "MOD-CC-001", "PLT-800", "91 s"],
    ],
    widths=[2.0, 3.3, 3.0, 1.2, 1.5], fontsize=8.5)
note("智能化实验室规划布局图",
     "因客户场地尺寸与隔断信息待补充，本版以工作站串联拓扑（见第 4、6 章平台俯视布局示意与产线串联流程）替代按实际墙体绘制的 CAD 规划图；正式布局 CAD 图于阶段一现场勘测后出具（见第 8 章交付物）。")
bullet("布局合理性：按前处理工序顺序线性排布，机械臂取放半径内完成工站间流转，操作便捷、流程合理；单平台占地 0.68 ㎡（PLT-800），空间利用率高。", size=9.5)
bullet("标准化预装与现场拼接：各平台 / 工站在工厂完成标准化预装与联调（FAT），现场仅做拼接与对接（SAT），大幅缩短现场建设周期（详见第 7 章实施阶段）。", size=9.5)

h(2, "3.3  配套基础设施改造设计")
note("说明",
     "下列改造要求由各模块卡 electrical_environment 与平台规格反推得出（数据可溯源），具体改造图（排风 / 气路 / 强弱电 / 给排水走向图）须在阶段一现场勘测后结合客户场地出具。")
table(
    ["系统", "改造设计要求", "来源依据"],
    [
        ["管道排风系统", "WS-02 离心、WS-04 浓缩、WS-05 过滤及方案二微波工站上方设万向抽气罩 / 防爆排风；提取涉乙腈，净化涉甲醇/乙酸乙酯/浓氨水/甲酸 + 45 ℃ 挥发，达标后高空排放。", "LX-001 requires_ventilation=true；1.html §10"],
        ["气路分布系统", "洁净干燥压缩空气 0.4–0.7 MPa（拆盖加液、SPE、过滤、涡旋气动压紧、机械臂）；真空泵：浓缩 −95 kPa、离心 −80 kPa。", "各模块 requires_compressed_air；NS-001 −95 kPa、LX-001 −80 kPa"],
        ["强弱电与网络", "主电 220 V 单相 + 控制电 24 V；浓缩 MOD-NS-001 另需 48 V；独立配电箱 + 漏电/接地保护；网络支持 MODBUS TCP 设备总线 + 软件 B/S 接入。", "各模块 voltage_v；NS-001 48V"],
        ["给排水 / 纯水管路", "一级纯水（GB/T 6682）供 SPE 活化/淋洗、均质/定容清洗；平台排废口接独立排废管路，有机废液与氨性废液分类收集至危废罐，不得直排市政。", "JZ-002/CS-004/SPE-004/NS-001 requires_water & requires_drain"],
    ],
    widths=[2.0, 6.4, 2.6], fontsize=8.5)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ====================================================================
# 4 功能平台逐一详细设计
# ====================================================================
h(1, "4  功能平台逐一详细设计")
para("对每一个已配置功能平台，逐一展开【平台组成 + 平台功能说明 + 关键技术指标表】三要素。指标均量化（数值 + 单位 + 精度）并标注通讯协议。", size=10)
note("通讯协议说明",
     "SmartLabOS 各模块经标准模块位接入平台统一设备总线，采用 MODBUS TCP 协议纳入智慧实验室软件统一调度（MOD-SPE-004 卡片明确『仅支持 MODBUS TCP』，新建 MOD-WB-001 URS 亦规定 MODBUS TCP）。"
     "对卡片未单列通讯字段的模块，统一标注为『MODBUS TCP（随平台总线统一调度）』，不作其他编造。")

# 各功能平台： (标题, 组成, 功能说明, 指标表rows)
# 指标表列： 模块ID | 设备名称 | 关键技术指标（数值+单位+精度）| 通讯协议
def platform_block(idx, title, members, func_text, spec_rows, layout_text):
    h(2, "4.%d  %s" % (idx, title))
    para("【平台组成】" + members, size=10, bold=False)
    para("【平台功能说明】" + func_text, size=10)
    table(["模块 ID", "设备名称", "关键技术指标（数值 + 单位 + 精度）", "通讯协议"],
          spec_rows, widths=[1.4, 2.3, 6.0, 1.3], fontsize=8.5,
          caption="关键技术指标表")
    para("【平台俯视布局示意】" + layout_text, size=9, color=GREY)

MODBUS = "MODBUS TCP\n（卡片明确）"
MODBUS_U = "MODBUS TCP\n（随平台总线\n统一调度）"

platform_block(1, "智能仓库 / 冷链存储平台（WS-06）",
    "MOD-CC-001 800 样品存储模块（占 3 模组位，仅兼容 PLT-800）。",
    "实现 2 mL 进样小瓶 0–10 ℃ 冷链存储与 XYZ 机械臂自动取放，作为待测液冷藏缓存与产线自动上下料缓冲站；单瓶取放 31 s（远低于 600 s），并行运行不参与产线瓶颈，保障痕量待测液不降解（对应挑战-2 痕量损失）。",
    [["MOD-CC-001", "800 样品存储模块",
      "存储工位 90 位（2 mL 西林瓶）；控温 0–10 ℃；XYZ 机械臂自动取放；单循环 31 s + 上下料 60 s；外形 850×800×600 mm，占 3 模组位；供电 220 V，额定 1.44 kVA / 峰值 2.4 kVA；压缩空气 1–4 bar；噪声 65 dB；维护周期 6 个月；交期 8 周", MODBUS_U]],
    "PLT-800 整机 90 工位冷藏存储，XYZ 机械臂自动取放。")

platform_block(2, "制样提取平台（WS-01）",
    "MOD-JZ-002 均质模块 + MOD-JY-010 拆装盖加液模块 + MOD-WX-004 单涡旋模块（PLT-800，模块位 3/3，节拍 566 s）。",
    "完成组织均质制样 → 拆盖加 Mcllvaine-Na₂EDTA 缓冲液 8 mL（方案二为乙腈-水 25 mL）→ 装盖涡旋混匀 1 min。高速均质刀头自清洗保证分散均一、规避基质分层与痕量取样偏差（对应挑战-2）。",
    [["MOD-JZ-002", "50ml 离心管均质模块",
      "处理容量 1–40 mL；均质转速 1000–30000 rpm；刀头自清洗；单循环 60 s + 上下料 60 s；外形 450×180×600 mm，占 1 位；供电 24 V，额定 0.2826 kVA / 峰值 0.471 kVA；需压缩空气 1–4 bar + 循环水；噪声 87 dB", MODBUS_U],
     ["MOD-JY-010", "50ml 离心管拆装盖加液模块",
      "加液量 1–50 mL；绝对精度 ±1%，重复精度 ±0.5%；定位精度 ±0.1 mm（Y 轴 100 mm / Z 轴 50 mm，50 mm/s）；含拆盖 165 / 装盖 166 / 加液 167（功能码 169）；单循环 106 s + 上下料 60 s；供电 24 V，0.228/0.38 kVA；压缩空气 1–4 bar", MODBUS_U],
     ["MOD-WX-004", "50ml 离心管单涡旋模块",
      "涡旋转速 1000–3000 rpm；振幅 ±2 mm；气动压紧；单次连续 ≤10 min，液量 ≤45 mL；单循环 220 s + 上下料 60 s；供电 24 V，0.228/0.38 kVA；压缩空气 1–4 bar", MODBUS_U]],
    "PLT-800 三模块位横排：均质 → 拆盖加液 → 涡旋。")

platform_block(3, "超声·冷冻离心平台（WS-02，方案一瓶颈站）",
    "MOD-CS-004 多通道超声模块（托盘位）+ MOD-LX-001 高速离心机模块（占 3 模组位，带 XYZ）（PLT-800，节拍 720 s）。",
    "超声辅助提取（标准要求 20 min，模块功能码 300 s，可参数化至 1200 s）→ 冷冻高速离心固液分离，取上清；残渣回 WS-01 补缓冲液重复提取一次后合并。多工位并行 + TEC 控温保障提取条件稳定。",
    [["MOD-CS-004", "50ml 离心管多通道超声模块",
      "多通道 16 工位并行；处理容量 1–50 mL；TEC 控温 25–60 ℃；冷却循环水；单循环 300 s（可参数化至 1200 s 满足超声 20 min）+ 上下料 60 s；供电 220 V，0.23/0.37 kVA；需循环水 + 排水；噪声 72 dB", MODBUS_U],
     ["MOD-LX-001", "50ml×4 孔位高速离心机模块（带 XYZ）",
      "4 孔位 ×50 mL；转速 1–8000 rpm；控温 2–7 ℃ 预冷；真空 −80 kPa；单循环 0–300 s（取上限 300 s）+ 上下料 60 s；外形 800×850×2050 mm，占 3 位、自重 201 kg；供电 24 V，额定 3.5 kVA / 峰值 4.8 kVA；需排风 + 真空泵 + 压缩机", MODBUS_U]],
    "PLT-800：离心机占 3 模组位（整机）+ 超声占托盘位。")

platform_block(4, "SPE 净化平台（WS-03，方案一专用）",
    "MOD-YY-004 全量液体转移模块 + MOD-SPE-004 6cc 自动上样 SPE 模块（PLT-800，模块位 2/3，节拍 520 s）。",
    "提取液全量合并并上样 → HLB 固相萃取柱『活化 - 上样 - 淋洗 - 洗脱』全自动八步可编程，注射泵微升级流速闭环规避柱干、流速不均与回收率波动（对应挑战-4、挑战-6）。标准 HLB 200 mg/6 mL 与模块 6cc HLB 柱型吻合。",
    [["MOD-YY-004", "50ml 全量液体转移模块",
      "处理容量 0–50 mL 全量旋转倒液转移；拆盖-倒液-装盖（功能码 155）；单循环 40 s + 上下料 60 s；外形 450×180×600 mm，占 1 位；供电 24 V，0.312/0.52 kVA；压缩空气 1–4 bar", MODBUS_U],
     ["MOD-SPE-004", "6cc 自动上样 SPE 模块",
      "6cc HLB 柱（换载具适配 3cc/10cc）；流速 0.1–15 mL/min，控制精度 ±1%（绝对/重复）；移液量 1–50 mL；活化-上样-淋洗-洗脱可编程；单循环 360 s + 上下料 60 s；供电 24 V，0.168/0.28 kVA；6 mm 气源 0.4–0.7 MPa；需水 + 排水", MODBUS]],
    "PLT-800：全量转移 / 上样 → 6cc HLB 自动 SPE。")

platform_block(5, "浓缩平台（WS-04）",
    "MOD-NS-001 10ml 大容量连续浓缩模块（占 2 模组位）（PLT-800，节拍 540 s）。",
    "洗脱液 / 提取液 45 ℃ 真空离心浓缩至干（方案二为正丙醇旋蒸近干、氮吹干）。真空降沸点 + 45–50 ℃ 控温 + 红外测温安全监控，规避热敏目标物分解与挥发损失（对应挑战-5）。",
    [["MOD-NS-001", "10ml 大容量连续浓缩模块",
      "接收瓶容量 1–4 mL（源瓶 50/100/200 mL）；真空 −95 kPa + 离心 5000–8000 rpm + 热风 45–50 ℃；红外测温超温保护；单循环约 480 s（工程估算，卡片 cycle_time 为空，蒸干终点须靠曲线判断）+ 上下料 60 s；外形 450×380×600 mm，占 2 位、自重 50 kg；供电 48 V，0.38/0.64 kVA；需真空泵 + 压缩机 + 水 + 排水", MODBUS_U]],
    "PLT-800：浓缩占 2 位，余 1 位可扩并行浓缩或升级 PLT-1200/1400。")

platform_block(6, "定容·过滤平台（WS-05）",
    "MOD-DR-003 注射器定容模块 + MOD-GL-001 2.5cc 注射器过滤模块（PLT-800，模块位 2/3，节拍 600 s，恰命中均衡基准）。",
    "复溶液 1.0 mL 精密定容溶解 → 0.22 µm 针式过滤至 2 mL 进样小瓶。注射器定容高精度 + 自动洗针规避交叉污染；正压预过滤除微粒，降基质效应（对应挑战-5、挑战-6）。",
    [["MOD-DR-003", "注射器定容模块（250ul/1000ul）",
      "定容量 0.25–1 mL（250/1000 µL 进样器）；绝对精度 ±0.6%，重复精度 ≤0.2%（RSD）；自动双工位洗针；单循环 270 s + 上下料 60 s；外形 450×180×600 mm，占 1 位；供电 24 V，0.138/0.23 kVA；需水 + 排水", MODBUS_U],
     ["MOD-GL-001", "2.5cc 注射器过滤模块",
      "滤液量 1–2.5 mL；25 mm 针式滤器，0.22 µm；过滤压力 1–6 bar 正压预过滤；单循环 210 s + 上下料 60 s；外形 450×180×600 mm，占 1 位；供电 24 V，0.098/0.164 kVA；压缩空气", MODBUS_U]],
    "PLT-800：注射器定容 → 0.22 µm 过滤。")

platform_block(7, "微波辅助萃取平台（方案二 WS-02，新建模块）",
    "MOD-WB-001 50ml 离心管微波辅助萃取模块（URS 拟建）+ MOD-LX-001 离心 + MOD-YY-004 移液（PLT-1200，5 模组位，节拍 580 s）。",
    "满足 GB/T 21316—2007 §7.1『乙腈-水 + 微波辐照辅助提取（700 W，30 s × 2 次）』工序。SmartLabOS 现有 97 模块及客户确认 11 模块清单中无微波能力，依指令输出新建模块 URS（详见配套 URS 文件）；过渡期可由超声 MOD-CS-004 作机理不同的替代，须经方法学验证。",
    [["MOD-WB-001\n（URS 拟建）", "50ml 离心管微波辅助萃取模块",
      "微波频率 2450 MHz；功率可调（含 700 W 档）；单次辐照 5–300 s 可设（默认 30 s）；红外/光纤测温超温断电（上限 ≤60 ℃）；微波泄漏 ≤5 mW/cm²，门连锁 + 急停 + 屏蔽腔；处理时间 120 s/循环（规范 60 s + 上下料 60 s）；占 1 位；供电 220 V + 24 V，建议预留 ≥1.5 kVA【状态：URS 拟建，待研发评审与样机验证】", MODBUS]],
    "PLT-1200（6 模组位）：微波（1 位）+ 离心（3 位）+ 移液（1 位）= 5 位。")

h(2, "4.8  自动化物流搬运")
para("由各平台内置 XYZ 机械臂 / 转运机构实现工站间样品自动流转（上下料估算 60 s/模块），对接后端 LC-MS/MS 进样器。MOD-CC-001 兼作进出料冷链缓存站。客户确认 11 模块清单内无独立行走式物流机器人模块，跨平台长距离搬运如有需求须另行确认配置（不编造模块）。", size=9.5)

h(2, "4.9  各模块关键技术指标核对清单（参数总表）")
table(
    ["模块 ID", "类别", "功能码", "量程/容量", "精度", "关键性能", "循环耗时(s)", "通讯"],
    [
        ["MOD-CC-001", "存储", "D210", "90 工位/2mL", "—", "0–10 ℃ 冷藏，XYZ 取放", "31", "MODBUS TCP*"],
        ["MOD-JZ-002", "均质", "115", "1–40 mL", "—", "1000–30000 rpm，刀头自清洗", "60", "MODBUS TCP*"],
        ["MOD-JY-010", "加液", "169", "1–50 mL", "±1% / ±0.5%", "拆盖-加液-装盖，定位 ±0.1mm", "106", "MODBUS TCP*"],
        ["MOD-WX-004", "涡旋", "223", "≤45 mL", "—", "1000–3000 rpm，振幅 ±2mm", "220", "MODBUS TCP*"],
        ["MOD-CS-004", "超声", "726", "1–50 mL", "—", "16 工位，TEC 25–60 ℃", "300", "MODBUS TCP*"],
        ["MOD-LX-001", "离心", "188", "4×50 mL", "—", "1–8000 rpm，2–7 ℃，真空 −80kPa", "0–300", "MODBUS TCP*"],
        ["MOD-YY-004", "移液", "155", "0–50 mL", "—", "全量旋转倒液转移", "40", "MODBUS TCP*"],
        ["MOD-SPE-004", "SPE", "121", "1–50 mL", "±1%", "6cc HLB，流速 0.1–15 mL/min", "360", "MODBUS TCP"],
        ["MOD-NS-001", "浓缩", "200", "1–4 mL", "—", "真空 −95kPa，45–50 ℃，5000–8000 rpm", "≈480 估", "MODBUS TCP*"],
        ["MOD-DR-003", "定容", "660/661", "0.25–1 mL", "±0.6% / ≤0.2%", "250/1000 µL，自动洗针", "270", "MODBUS TCP*"],
        ["MOD-GL-001", "过滤", "637", "1–2.5 mL", "—", "0.22 µm，正压 1–6 bar", "210", "MODBUS TCP*"],
        ["MOD-WB-001", "微波(拟建)", "WB10", "5–300 s", "—", "2450 MHz，700 W，泄漏 ≤5mW/cm²", "120", "MODBUS TCP"],
    ],
    widths=[1.5, 1.0, 0.9, 1.2, 1.3, 3.2, 1.1, 1.3], fontsize=8)
para("注：* 表示卡片未单列通讯字段，随平台总线统一 MODBUS TCP 调度；MOD-LX-001 循环 0–300 s 取上限 300 s；MOD-NS-001 卡片 cycle_time 为空，按同类工艺工程估算 480 s；MOD-WB-001 为 URS 拟建模块，参数待样机验证。", size=8.5, color=GREY)

h(2, "4.10  辅助 / 可选设备（○ 按需）")
table(
    ["设备", "用途", "建议规格", "备注"],
    [
        ["纯水制备系统", "SPE 活化/淋洗、均质/定容清洗用水", "符合 GB/T 6682 一级纯水；产水能力按峰值用量核算", "按需选配；预算请联系销售团队"],
        ["UPS 不间断电源", "保障离心/浓缩等关键工序断电安全", "建议每平台 3 kVA 在线式", "按需选配（1.html §10）"],
        ["外置高速冷冻离心", "满足方案一 10000/14000 r/min 偏差（见 6 章/偏差）", "≥14000 r/min，控温至 −2 ℃", "如客户要求严格符合标准须授权增配（清单外）"],
    ],
    widths=[2.2, 4.0, 3.3, 2.5], fontsize=8.5)
para("以上辅助设备为按需裁剪项；如本项目暂不纳入，相关合规偏差处置须按第 6 章偏差说明与客户确认。涉及价格、报价信息，请联系销售团队获取报价。", size=9, color=GREY)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ====================================================================
# 5 软件系统详细设计
# ====================================================================
h(1, "5  软件系统详细设计")
note("说明",
     "客户软件需求为『智慧实验室软件』。该软件为 SmartLabOS 系统软件层，在 references/01-modules 硬件模块库中无对应硬件卡片，由 SmartLabOS 软件平台单独配置统一调度。本章按详细设计模板列出标准能力项；与客户 LIMS 的具体接口与数据流向须在阶段一对接确认。")

h(2, "5.1  中央控制系统")
para("技术架构：采用 B/S 架构、分布式部署，分『设备执行层 — 调度控制层 — 数据服务层 — 应用展示层』四层；设备层经 MODBUS TCP 总线连接各模块/平台，应用层经浏览器访问，支持集中-分散控制。", size=9.5)
para("功能模块：基础数据、销售管理、样品管理、仓库管理、实验流程管理、实验调度、设备管理与自动化执行、数字化大屏。", size=9.5)
para("自动化功能：", size=10, bold=True)
for t in [
    "实验室调度控制与单机设备执行：按检测标准库/实验流程库下发工序模板，控制各模块时序与并行；",
    "仓库管理与物料调度：MOD-CC-001 冷链库存监测、补料提醒；",
    "样品全生命周期管理：采样前/中/后 + 申请-审批-确认-记录闭环；",
    "设备台账与运维：保养提醒、故障报警、正反向追溯。",
]:
    bullet(t, size=9.5)
para("数据管理功能：实验数据抓取与传输、LIMS 双向对接、通风/有害气体监测联动、多通道实时采集与校验、数据清洗、数据审查与质控校验、实时存储与多维查询。", size=9.5)
para("其他：智能监测分析与预警；报表与统计（设备使用、检测次数、试剂用量、异常判定、原始记录/报告导出）；报告定制与一键生成（模板定制、多级审核）；数据导出（Word/Excel/PDF）；角色与权限分级管理；报警管理（实时/历史、规则配置、多通道通知）；数据安全（冗余容错、传输存储加密、权限隔离）。", size=9.5)

h(2, "5.2  数字孪生 / 可视化系统（○ 按需）")
para("可选配数字孪生可视化：中央控制可视化、试剂耗材管理、物流与设备运行监控；五大功能为实时数据展示与监控、实验过程可视化（视频流/多视角/录制）、数据分析与决策支持、预警与应急响应、成果展示与参观交互。本项目可按需裁剪，如纳入须单独确认范围。", size=9.5)

h(2, "5.3  软件系统技术指标")
table(
    ["指标类别", "指标要求"],
    [
        ["开放与集成", "提供 RESTful API；HTTPS 安全传输；设备侧通用 MODBUS TCP；支持与客户 LIMS / 工厂信息化双向对接（接口与数据流向阶段一确认）"],
        ["安全性", "容错、可扩充、操作日志完备、关键功能冗余切换"],
        ["可靠性", "成熟案例支撑；支持 7×24 运行；集中-分散控制；年故障累计时间上限按合同 SLA 约定"],
        ["可扩展性", "硬件可扩容（新增模块/平台免改架构）；软件免定制配置适配；升级简便"],
        ["日志记录", "记录操作人/时间/动作/结果等要素；保存时长可设；到期自动清理"],
    ],
    widths=[2.2, 8.8], fontsize=9)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ====================================================================
# 6 自动化产能效率
# ====================================================================
h(1, "6  自动化产能效率")
para("效率设计逻辑：多模块并行（离心 4 孔、超声 16 工位、SPE 可并行）+ 全流程联动 + 工站节拍匹配（按『每平台 ≈ 600 s』均衡拆分），稳态流水生产时各工站重叠运行，整体通量由瓶颈站决定。", size=9.5)

h(2, "6.1  方案一逐平台耗时测算（GB 31658.17—2021）")
para("口径：每平台处理耗时 = Σ（功能码 cycle_time_sec + 自动上下料 60 s/模块）。基质：牛/羊/猪/鸡 肌肉·肝·肾；目标物 41 种。", size=9, color=GREY)
table(
    ["工作站（平台）", "模块（功能码耗时 + 上下料）", "平台处理耗时", "相对 600 s 基准"],
    [
        ["WS-01 制样提取", "JZ-002(60+60) + JY-010(106+60) + WX-004(220+60)", "566 s", "−34 s"],
        ["WS-02 超声·离心", "CS-004(300+60) + LX-001(300+60)", "720 s", "+120 s（瓶颈）"],
        ["WS-03 转移·SPE", "YY-004(40+60) + SPE-004(360+60)", "520 s", "−80 s"],
        ["WS-04 浓缩", "NS-001(480※+60)", "540 s", "−60 s"],
        ["WS-05 定容·过滤", "DR-003(270+60) + GL-001(210+60)", "600 s", "±0 s（命中）"],
        ["WS-06 冷链存储", "CC-001(31+60)", "91 s", "并行缓存，不计瓶颈"],
        ["解决方案总节拍（串行单样）", "Σ 五处理站 + 缓存站", "3037 s ≈ 50.6 min/样", "5 处理站均值 589 s"],
    ],
    widths=[2.6, 5.0, 2.0, 1.4], fontsize=8.5)

h(2, "6.2  方案二逐平台耗时测算（GB/T 21316—2007）")
para("口径同上。基质：肝/肾/肌肉/水产/牛奶；目标物 23 种磺胺类。含新建微波模块 MOD-WB-001（120 s/循环）。", size=9, color=GREY)
table(
    ["工作站（平台）", "模块组成", "平台处理耗时", "说明"],
    [
        ["WS-01 研磨提取（PLT-800）", "JZ-002 + JY-010 + WX-004", "566 s", "C₁₈ 研磨以均质近似"],
        ["WS-02 微波萃取·离心（PLT-1200）", "MOD-WB-001(新) + LX-001 + YY-004", "580 s", "微波 120 s + 离心 360 s + 移液 100 s"],
        ["WS-03 浓缩（PLT-800）", "NS-001", "540 s", "正丙醇 45 ℃ 旋蒸/氮吹干"],
        ["WS-04 复溶定容（PLT-800）", "DR-003", "330 s", "乙腈-水 1 mL 复溶定容"],
        ["WS-05 超声·过滤（PLT-800）", "CS-004 + GL-001", "630 s", "超声 30 s 溶解 + 0.22 µm 过滤（瓶颈）"],
        ["WS-06 冷链存储（PLT-800）", "CC-001", "91 s", "并行缓存"],
        ["解决方案总节拍（串行单样）", "Σ 各站", "2737 s ≈ 45.6 min/样", "PLT-800×5 + PLT-1200×1"],
    ],
    widths=[2.8, 3.6, 2.0, 2.6], fontsize=8.5)

h(2, "6.3  高通量与稳态通量")
bullet("稳态流水：方案一瓶颈站 WS-02 = 720 s，理想稳态约 5 样/小时；方案二瓶颈站 WS-05 = 630 s，理想稳态约 5.7 样/小时（串行单样节拍仅作单样全流程参考）。", size=9.5)
bullet("并行扩容：离心 4 孔、超声 16 工位天然多工位并行；瓶颈站可在平台余位扩并行模块或升级 PLT-1200/1400，进一步提升稳态通量。", size=9.5)
bullet("专项/痕量检测（如硝基呋喃含恒温衍生等复杂环节）：客户确认 11 模块清单中无恒温衍生模块，如该检项纳入须另行确认配置；本方案现有效率测算覆盖四环素/磺胺/喹诺酮（方案一）与磺胺（方案二）两条已确认产线。", size=9.5)

h(2, "6.4  效率提升对比（vs 人工）")
note("说明",
     "人工基线（人力 N→M、错误率 X%→Y%）需以客户现场实测数据标定，本版不编造具体数值，给出可量化的自动化设计收益方向；正式立项后以现场标定补充对比表。")
for t in [
    "一致性：参数固化（转速/流速/温度/时间/体积）+ 全自动上下料，消除人工手法差异，提升批间重复性（针对挑战-1/-2）；",
    "回收率：SPE 流速闭环（±1%）、定容精度 ±0.6%/RSD ≤0.2%、真空控温浓缩防分解，系统性降低损失与基质效应（针对挑战-3/-4/-5/-6）；",
    "效率与无人化：单样全流程自动化串行节拍方案一 50.6 min、方案二 45.6 min，稳态流水重叠运行支持高通量筛查；冷链缓存站支持无人值守续跑。",
]:
    bullet(t, size=9.5)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ====================================================================
# 7 项目实施与管理方案
# ====================================================================
h(1, "7  项目实施与管理方案")

h(2, "7.1  周期与组织")
para("总建设周期（估算）：平台/模块标准交付约 8 周（卡片 lead_time_weeks），工站集成约 12 周；叠加现场施工、安装拼接、调试培训与验收，总周期约 120–150 日历天（以阶段一细化与合同为准）。", size=9.5)
table(
    ["角色", "职责"],
    [
        ["乙方项目经理", "总体计划、资源协调、进度与风险管控、对客户接口"],
        ["乙方技术 / 工艺团队", "工艺细化、模块集成、流程模板与方法学验证支持"],
        ["乙方质控团队", "FAT/SAT 质控、文档归档、验收交付"],
        ["甲方对接小组", "需求确认、场地与公用工程配合、方法学样品与验收参与"],
    ],
    widths=[2.6, 8.4], fontsize=9)

h(2, "7.2  实施阶段")
table(
    ["阶段", "工期(估)", "核心任务", "交付物", "验收标准 / 管控要点"],
    [
        ["一、调研细化与方案确认", "2–3 周", "现场勘测、场地/公用工程确认、偏差处置确认、方案与流程模板细化", "细化设计方案、布局 CAD 图、确认清单", "客户书面确认；偏差项闭环"],
        ["二、模块生产与工厂预装调试", "8 周", "模块生产、平台标准化预装、工厂联调（FAT）", "FAT 报告、设备合格证", "FAT 通过；节拍/精度达标"],
        ["三、现场基础施工与复核", "2–3 周", "排风/气路/强弱电/给排水改造、承重与层高复核", "施工记录、复核报告", "公用工程达标；与设备接口匹配"],
        ["四、现场安装与模块拼接", "2 周", "平台进场、拼接、对接 LC-MS/MS", "安装/拼接验收单", "拼接到位；通电通气通水正常"],
        ["五、系统调试与人员培训", "3–4 周", "联调（SAT）、方法学验证、操作/运维培训", "SAT 报告、培训记录、方法学验证报告", "回收率/RSD 满足国标；人员考核合格"],
        ["六、项目验收与交付归档", "1–2 周", "终验、资料归档、移交", "验收报告、设备移交清单、全套文档", "终验通过；文档齐全可追溯"],
    ],
    widths=[2.2, 1.0, 3.0, 2.4, 2.4], fontsize=8)
note("总工期甘特图",
     "六阶段为部分串行、部分可重叠（阶段二工厂预装与阶段三现场施工可并行）。甘特图于阶段一按客户确认的开工日与场地条件出具（见第 8 章交付物）。")

h(2, "7.3  项目管理措施")
for t in [
    "进度管理：三级管控（日跟踪 / 周例会 / 里程碑评审）+ 赶工预案；",
    "质量管理：全流程质控追溯，FAT/SAT 质控节点 + 归档；",
    "沟通管理：固定例会机制 + 专属沟通群 + 问题清单闭环；",
    "安全管理：用电/动火/接地/急停规范，权限分级，有机溶剂与微波安全防护；",
    "变更管理：书面申请 → 影响评估 → 签字确认。",
]:
    bullet(t, size=9.5)

h(2, "7.4  保障措施")
for t in [
    "资源保障：资深工程师 + 专用设备 + 原厂部件，甲方配合到位；",
    "技术保障：攻坚小组 + 7×24 远程支持 + 现场响应时限（按合同 SLA）；",
    "售后前置：售后工程师参与调试，提前熟悉系统；",
    "文档保障：全流程归档可追溯。",
]:
    bullet(t, size=9.5)

# ====================================================================
# 8 交付物清单
# ====================================================================
h(1, "8  交付物清单")
table(
    ["类别", "交付物"],
    [
        ["正式技术文档", "详细设计方案、各类设计图（布局/排风/气路/强弱电/给排水）"],
        ["工程交付物", "布局 CAD 图、工厂(FAT)/现场(SAT)调试报告、设备合格证、安装/拼接验收单、验收报告、设备移交清单"],
        ["资料交付", "操作手册、维护手册、技术资料（纸质 + 电子）、售后服务协议"],
        ["方法学", "方法学验证报告（回收率/RSD 满足 GB 31658.17—2021 / GB/T 21316—2007）"],
    ],
    widths=[2.2, 8.8], fontsize=9)

# ====================================================================
# 附录：偏差与待确认事项
# ====================================================================
h(1, "附录 A  工艺偏差与待客户确认事项（诚实标注）")
para("以下偏差源于客户已确认 11 模块清单的能力边界，须与客户确认处置方式（不在清单外自行替换模块）：", size=9.5)
table(
    ["编号", "偏差 / 待确认项", "标准要求", "本方案能力 / 处置建议"],
    [
        ["D-1", "方案一提取离心（GB 31658.17 §8.1）", "−2 ℃ / 10000 r/min", "MOD-LX-001 为 2–7 ℃ / ≤8000 rpm，无法严格达到；建议①方法学验证 8000 rpm·2 ℃ 等效，或②清单外增配外置高速冷冻离心（须授权）"],
        ["D-2", "方案一复溶离心", "14000 r/min", "清单内仅 MOD-LX-001(≤8000 rpm)；以 0.22 µm 过滤(MOD-GL-001)替代去微粒，必要时叠加 8000 rpm 预沉降；严格符合须增配高速离心"],
        ["D-3", "方案一超声时长（§8.1）", "超声 20 min（1200 s）", "MOD-CS-004 功能码 300 s，可参数化配置至 1200 s；耗时测算按 300 s 计，配置后 WS-02 节拍相应延长，可经多通道并行维持均衡"],
        ["D-4", "方案二微波辐照（GB/T 21316 §7.1）", "微波 700 W，30 s × 2 次", "清单内无微波模块，输出新建 MOD-WB-001 URS（待立项/样机验证）；过渡期可用超声替代（机理不同，须方法学验证）"],
        ["D-5", "方案二 C₁₈ 研磨 / 液液分配净化", "MSPD 研磨 / 正己烷液液分配", "以均质 MOD-JZ-002 近似分散，液液分配复用涡旋/离心/移液组合等效，须方法学验证"],
        ["D-6", "称样", "1 g / 2 g（0.01 g 天平）", "清单内无称量模块，按惯例人工称样后上样"],
        ["D-7", "智慧实验室软件", "系统软件层", "硬件模块库无对应卡片，由 SmartLabOS 软件平台单独配置；与 LIMS 接口阶段一确认"],
    ],
    widths=[0.7, 2.8, 2.3, 5.2], fontsize=8)
note("数据来源与合规声明",
     "本方案全部模块、平台技术参数均引自 SmartLabOS 知识库 references/01-modules、02-platforms、04-solutions，并严格限定在客户已确认的 11 个模块清单内选型；缺口（微波）以新建模块 URS 形式补充，未编造既有模块 ID 与参数。涉及价格、报价信息，请联系销售团队获取报价。")

# ---------- 页脚页码 ----------
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("百泉聚兴（北京）科技有限公司 · SmartLabOS · 测试项目详细设计方案    第 ")
r.font.size = Pt(8); r.font.color.rgb = GREY; set_eastasia(r)
add_field(fp, "PAGE")
r2 = fp.add_run(" 页 / 共 "); r2.font.size = Pt(8); r2.font.color.rgb = GREY; set_eastasia(r2)
add_field(fp, "NUMPAGES")
r3 = fp.add_run(" 页"); r3.font.size = Pt(8); r3.font.color.rgb = GREY; set_eastasia(r3)

doc.save(OUT)
print("SAVED:", OUT)
print("EXISTS:", os.path.exists(OUT), "SIZE:", os.path.getsize(OUT))
